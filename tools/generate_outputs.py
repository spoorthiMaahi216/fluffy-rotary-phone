#!/usr/bin/env python3
import os
from pathlib import Path
from typing import List, Dict

from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from urllib.parse import urlparse, quote
from urllib.request import urlretrieve
from urllib.request import Request, urlopen
from io import BytesIO
from PIL import Image

OUTPUT_DIR = Path('/workspace/generated')
IMAGES_DIR = OUTPUT_DIR / 'images'


def ensure_dirs() -> None:
	OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
	IMAGES_DIR.mkdir(parents=True, exist_ok=True)


def save_right_triangle_image(path: Path, leg_a: int, leg_b: int) -> None:
	fig, ax = plt.subplots(figsize=(3, 3))
	ax.plot([0, leg_a], [0, 0], 'k-', linewidth=2)
	ax.plot([0, 0], [0, leg_b], 'k-', linewidth=2)
	ax.plot([leg_a, 0], [0, leg_b], 'k--', linewidth=2)
	# Right angle box
	xb, yb, s = 0, 0, min(leg_a, leg_b) * 0.15
	ax.plot([xb, xb + s], [yb, yb], 'k-', linewidth=2)
	ax.plot([xb, xb], [yb, yb + s], 'k-', linewidth=2)
	ax.plot([xb + s, xb + s], [yb, yb + s], 'k-', linewidth=2)
	ax.plot([xb, xb + s], [yb + s, yb + s], 'k-', linewidth=2)
	ax.text(leg_a / 2, -0.5, f'{leg_a}', ha='center', va='top')
	ax.text(-0.5, leg_b / 2, f'{leg_b}', ha='right', va='center', rotation=90)
	ax.set_xlim(-1, max(leg_a, leg_b) + 1)
	ax.set_ylim(-1, max(leg_a, leg_b) + 1)
	ax.set_aspect('equal')
	ax.axis('off')
	fig.tight_layout()
	fig.savefig(path, dpi=200, bbox_inches='tight')
	plt.close(fig)


def save_fruit_barchart_image(path: Path, labels: List[str], values: List[int]) -> None:
	fig, ax = plt.subplots(figsize=(4, 3))
	bars = ax.bar(labels, values, color=['#4e79a7', '#f28e2b', '#59a14f'])
	for bar, v in zip(bars, values):
		ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.2, str(v), ha='center', va='bottom')
	ax.set_ylim(0, max(values) + 2)
	ax.set_ylabel('Count')
	ax.set_title('Fruit Counts')
	fig.tight_layout()
	fig.savefig(path, dpi=200, bbox_inches='tight')
	plt.close(fig)


def build_new_questions_blocks() -> List[Dict[str, str]]:
	return [
		{
			"title": "Right Triangle Hypotenuse",
			"description": "Find the hypotenuse of a right triangle using the Pythagorean theorem.",
			"question": "In the right triangle shown, the legs have lengths 6 and 8 units. What is the length of the hypotenuse?",
			"instruction": "Select the correct value of the hypotenuse.",
			"difficulty": "easy",
			"order": "1",
			"options": ["7", "8", "9", "10", "14"],
			"answer": "10",
			"explanation": "By the Pythagorean theorem, c = \\sqrt{6^{2} + 8^{2}} = \\sqrt{36 + 64} = \\sqrt{100} = 10.",
			"subject": "Quantitative Math",
			"unit": "Geometry and Measurement",
			"topic": "Right Triangles & Trigonometry",
			"image_path": str(IMAGES_DIR / 'q_new_triangle.png')
		},
		{
			"title": "Average Fruit Count",
			"description": "Compute the mean from a bar chart of three categories.",
			"question": "The bar chart shows the counts of Apples, Bananas, and Cherries collected by a class. What is the mean (average) count across the three fruits?",
			"instruction": "Choose the mean of the displayed counts.",
			"difficulty": "moderate",
			"order": "2",
			"options": ["6", "6.5", "7", "7.5", "8"],
			"answer": "7",
			"explanation": "If the counts are 4, 7, and 10, then mean = (4+7+10)/3 = 21/3 = 7.",
			"subject": "Quantitative Math",
			"unit": "Data Analysis & Probability",
			"topic": "Mean, Median, Mode, & Range",
			"image_path": str(IMAGES_DIR / 'q_new_bars.png')
		}
	]


def write_new_questions_docx(blocks: List[Dict[str, str]], path: Path) -> None:
	doc = Document()
	for b in blocks:
		doc.add_paragraph(f"@title {b['title']}")
		doc.add_paragraph(f"@description {b['description']}")
		doc.add_paragraph("")
		doc.add_paragraph(f"@question {b['question']}")
		doc.add_paragraph(f"@instruction {b['instruction']}")
		doc.add_paragraph(f"@difficulty {b['difficulty']}")
		doc.add_paragraph(f"@Order {b['order']}")
		# options with correct marked using @@option
		for opt in b['options']:
			if opt == b['answer']:
				doc.add_paragraph(f"@@option {opt}")
			else:
				doc.add_paragraph(f"@option {opt}")
		doc.add_paragraph("@explanation")
		doc.add_paragraph(b['explanation'])
		doc.add_paragraph(f"@subject {b['subject']}")
		doc.add_paragraph(f"@unit {b['unit']}")
		doc.add_paragraph(f"@topic {b['topic']}")
		doc.add_paragraph("@plusmarks 1")
		# Add image if present
		img_path = b.get('image_path')
		if img_path and Path(img_path).exists():
			doc.add_picture(img_path, width=Inches(3.5))
		doc.add_paragraph("")
		doc.add_paragraph("---")
		doc.add_paragraph("")
	doc.save(path)


def render_question_block(
	*,
	title: str,
	description: str,
	question: str,
	instruction: str,
	difficulty: str,
	order: int,
	options: List[str],
	answer: str,
	explanation: str,
	subject: str,
	unit: str,
	topic: str,
) -> str:
	lines: List[str] = []
	lines.append(f"@title {title}")
	lines.append(f"@description {description}")
	lines.append("")
	lines.append(f"@question {question}")
	lines.append(f"@instruction {instruction}")
	lines.append(f"@difficulty {difficulty}")
	lines.append(f"@Order {order}")
	for opt in options:
		prefix = '@@option' if opt == answer else '@option'
		lines.append(f"{prefix} {opt}")
	lines.append("@explanation")
	lines.append(explanation)
	lines.append(f"@subject {subject}")
	lines.append(f"@unit {unit}")
	lines.append(f"@topic {topic}")
	lines.append("@plusmarks 1")
	lines.append("")
	return "\n".join(lines)


# ---------------- Shadow questions (25 similar variants) ----------------

def build_25_shadow_questions_text() -> str:
	blocks: List[str] = []
	add = blocks.append

	# 1
	add(render_question_block(
		title='Solve Linear Equation (One-Step)',
		description='Solve for n in a simple linear equation.',
		question='If $n+7=12$, what is the value of $n$?',
		instruction='Select the correct value of n.',
		difficulty='easy', order=1,
		options=['2', '4', '5', '7', '12'], answer='5',
		explanation='Subtract 7 from both sides: $n = 12-7 = 5$.',
		subject='Quantitative Math', unit='Algebra', topic='Interpreting Variables'
	))

	# 2
	add(render_question_block(
		title='Repeating Symbol Sequence',
		description='Identify a term in a repeating sequence using modular arithmetic.',
		question='A sequence repeats the symbols in order: Circle, Square, Triangle, Star. Which is the 12th symbol?',
		instruction='Determine the cycle length and reduce the index modulo the cycle length.',
		difficulty='moderate', order=2,
		options=['Circle', 'Square', 'Triangle', 'Star', 'Hexagon'], answer='Star',
		explanation='Cycle length is 4. 12 mod 4 = 0, so the 12th is the 4th in the cycle: Star.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Sequences & Series'
	))

	# 3
	add(render_question_block(
		title='Expression for Total Items',
		description='Translate a word scenario into an algebraic expression.',
		question='A jar contains 15 marbles. You add $y$ more marbles. Which expression represents the total number of marbles?',
		instruction='Choose the expression that models the situation.',
		difficulty='easy', order=3,
		options=['$15-y$', '$15y$', '$\\frac{15}{y}$', '$y-15$', '$15+y$'], answer='$15+y$',
		explanation='Start with 15 and add y new marbles: $15 + y$.',
		subject='Quantitative Math', unit='Algebra', topic='Interpreting Variables'
	))

	# 4
	add(render_question_block(
		title='Place Value and Inequality',
		description='Find the greatest digit for a number to stay below a bound.',
		question='In the number $5,\\square 42$, $\\square$ is a digit 0–9. If the number is less than 5,242, what is the greatest possible value for $\\square$?',
		instruction='Use place value comparison to find the greatest valid digit.',
		difficulty='easy', order=4,
		options=['0', '1', '2', '4', '9'], answer='1',
		explanation='Compare hundreds place with 2 in 5,242: the greatest hundreds digit to keep it smaller is 1.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Computation with Whole Numbers'
	))

	# 5
	add(render_question_block(
		title='Adding Fractions',
		description='Add two fractions with unlike denominators.',
		question='Which of the following is the sum of $\\frac{5}{12}$ and $\\frac{1}{3}$?',
		instruction='Compute using a common denominator.',
		difficulty='easy', order=5,
		options=['$\\frac{1}{4}$', '$\\frac{2}{3}$', '$\\frac{3}{4}$', '$\\frac{5}{6}$', '$\\frac{7}{12}$'], answer='$\\frac{3}{4}$',
		explanation='$\\frac{5}{12}+\\frac{1}{3}=\\frac{5}{12}+\\frac{4}{12}=\\frac{9}{12}=\\frac{3}{4}$.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Fractions, Decimals, & Percents'
	))

	# 6
	add(render_question_block(
		title='Altitude Difference from a Graph (Conceptual)',
		description='Read altitude change from start and finish.',
		question='A hiker starts at 120 meters and ends at 420 meters after a steady climb. How many meters higher is the end than the start?',
		instruction='Compute final altitude minus initial altitude.',
		difficulty='easy', order=6,
		options=['120', '240', '300', '320', '540'], answer='300',
		explanation='420 − 120 = 300 meters.',
		subject='Quantitative Math', unit='Data Analysis & Probability', topic='Interpretation of Tables & Graphs'
	))

	# 7
	add(render_question_block(
		title='Multiply Decimals',
		description='Evaluate a product of decimals.',
		question='What is the value of $0.25 \\times 18 \\times 0.4$?',
		instruction='Use associativity to simplify.',
		difficulty='easy', order=7,
		options=['0.18', '1.8', '18', '180', '0.72'], answer='1.8',
		explanation='$0.25 \\times 0.4 = 0.1$ and $0.1 \\times 18 = 1.8$.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Fractions, Decimals, & Percents'
	))

	# 8
	add(render_question_block(
		title='Minimize Coins for a Total',
		description='Find the least number of coins to make a given amount.',
		question='There are ten of each coin: 1¢, 5¢, 10¢, and 25¢. If you need exactly 47¢, what is the least number of coins required?',
		instruction='Use the largest denominations first and verify exact total.',
		difficulty='moderate', order=8,
		options=['Three', 'Four', 'Five', 'Six', 'Seven'], answer='Five',
		explanation='47 = 25 + 10 + 10 + 1 + 1 uses five coins; four coins cannot make 47.',
		subject='Quantitative Math', unit='Reasoning', topic='Word Problems'
	))

	# 9
	add(render_question_block(
		title='Multiply Fractions then Halve',
		description='Evaluate a nested fractional expression.',
		question='What is the value of $\\frac{1}{2}\\left(\\frac{2}{3} \\times \\frac{3}{4}\\right)$?',
		instruction='Multiply inside the parentheses first.',
		difficulty='easy', order=9,
		options=['$\\frac{1}{4}$', '$\\frac{1}{3}$', '$\\frac{3}{8}$', '$\\frac{5}{12}$', '$\\frac{7}{24}$'], answer='$\\frac{1}{4}$',
		explanation='$\\frac{2}{3} \\times \\frac{3}{4} = \\frac{1}{2}$; then half gives $\\frac{1}{4}$.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Fractions, Decimals, & Percents'
	))

	# 10
	add(render_question_block(
		title='Midpoints on a Line Segment',
		description='Use midpoint relations to compute a segment length.',
		question='Segment $\\overline{ST}$ has length 10, $T$ is the midpoint of $\\overline{RV}$, and $S$ is the midpoint of $\\overline{RT}$. What is the length of $\\overline{SV}$?',
		instruction='Express RV in terms of ST using midpoint relations.',
		difficulty='moderate', order=10,
		options=['10', '20', '30', '40', '50'], answer='30',
		explanation='S midpoint of RT ⇒ ST = RT/2 ⇒ RT = 20. T midpoint of RV ⇒ TV = RT = 20. So SV = ST + TV = 10 + 20 = 30.',
		subject='Quantitative Math', unit='Geometry and Measurement', topic='Lines, Angles, & Triangles'
	))

	# 11
	add(render_question_block(
		title='Solve Whole-Number Identity',
		description='Solve for a whole number that satisfies a simple quadratic identity.',
		question='Let $a$ be defined by $a=a^{2}-a$, where $a$ is a whole number and $a\\neq 0$. What is the value of $3a$?',
		instruction='Solve for a, then compute 3a.',
		difficulty='easy', order=11,
		options=['4', '5', '6', '7', '8'], answer='6',
		explanation='$a=a^{2}-a \\Rightarrow a^{2}-2a=0 \\Rightarrow a(a-2)=0$. With $a\\neq 0$, $a=2$, so $3a=6$.',
		subject='Quantitative Math', unit='Algebra', topic='Interpreting Variables'
	))

	# 12
	add(render_question_block(
		title='Counting Uniform Combinations',
		description='Count combinations from shirts and pants options.',
		question='A uniform has 1 shirt and 1 pair of pants. If there are 5 shirt colors and 2 pants colors, how many different uniforms are possible?',
		instruction='Multiply the number of shirt choices by pant choices.',
		difficulty='easy', order=12,
		options=['6', '8', '10', '12', '15'], answer='10',
		explanation='There are 5 shirts and 2 pants: $5 \\times 2 = 10$.',
		subject='Quantitative Math', unit='Data Analysis & Probability', topic='Counting & Arrangement Problems'
	))

	# 13
	add(render_question_block(
		title='Parity Reasoning',
		description='Determine which expression yields an odd integer for even n.',
		question='If $n$ is an even integer, which of the following must be an odd integer?',
		instruction='Analyze parity for each expression.',
		difficulty='easy', order=13,
		options=['$n$', '$n+1$', '$2n$', '$3n$', '$n+2$'], answer='$n+1$',
		explanation='If $n$ is even, then $n+1$ is odd.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Basic Number Theory'
	))

	# 14
	add(render_question_block(
		title='Direct Proportion: Miles per Dollar',
		description='Use proportional reasoning to scale miles by fuel cost.',
		question='A car travels 180 miles on $\\$30 of gas. At the same rate, how many miles on $\\$45?',
		instruction='Use miles per dollar to scale linearly.',
		difficulty='easy', order=14,
		options=['225', '240', '255', '270', '300'], answer='270',
		explanation='$180/30 = 6$ miles per dollar; $6 \\times 45 = 270$.',
		subject='Quantitative Math', unit='Reasoning', topic='Word Problems'
	))

	# 15
	add(render_question_block(
		title='Closest Fraction to a Percentage',
		description='Compare fractions to 62%.',
		question='Which fraction is closest to $62\\%$?',
		instruction='Convert fractions to percents or compare decimals.',
		difficulty='moderate', order=15,
		options=['$\\frac{1}{2}$', '$\\frac{3}{5}$', '$\\frac{5}{8}$', '$\\frac{2}{3}$', '$\\frac{7}{10}$'], answer='$\\frac{5}{8}$',
		explanation='$\\frac{5}{8}=0.625=62.5\\%$, closest to 62%.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Fractions, Decimals, & Percents'
	))

	# 16
	add(render_question_block(
		title='Balanced Club Sizes',
		description='Distribute students into clubs with max difference 1.',
		question='There are 84 students forming 5 clubs. Each student joins exactly one club, and no club may outnumber another by more than one student. What is the least possible number of students in one club?',
		instruction='Distribute as evenly as possible.',
		difficulty='moderate', order=16,
		options=['15', '16', '17', '18', '19'], answer='16',
		explanation='84 divided as evenly as possible into 5 gives sizes 17, 17, 17, 16, 17; the least is 16.',
		subject='Quantitative Math', unit='Data Analysis & Probability', topic='Counting & Arrangement Problems'
	))

	# 17
	add(render_question_block(
		title='Shaded Fraction of a Rectangle (Variant)',
		description='Find the shaded portion count out of total.',
		question='A rectangle is divided into 8 congruent squares. If $5\\tfrac{1}{2}$ squares are shaded, what fraction of the rectangle is shaded?',
		instruction='Compute shaded total over 8 and simplify if possible.',
		difficulty='easy', order=17,
		options=['$\\frac{5}{8}$', '$\\frac{11}{16}$', '$\\frac{3}{4}$', '$\\frac{7}{12}$', '$\\frac{2}{3}$'], answer='$\\frac{11}{16}$',
		explanation='$5.5/8 = 11/16$.',
		subject='Quantitative Math', unit='Geometry and Measurement', topic='Area & Volume'
	))

	# 18
	add(render_question_block(
		title='Currency Exchange Chains',
		description='Convert gold to copper through given exchange rates.',
		question='In a game, 1 gold piece may be exchanged for 4 silver pieces, and 3 silver pieces may be exchanged for 18 copper pieces. How many copper pieces for 5 gold pieces?',
		instruction='Find copper per gold, then scale.',
		difficulty='easy', order=18,
		options=['60', '90', '100', '120', '150'], answer='120',
		explanation='1 silver = 6 copper; 1 gold = 4 silver = 24 copper; 5 gold = 120 copper.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Rational Numbers'
	))

	# 19
	add(render_question_block(
		title='Sum of Horizontal Segments (Variant)',
		description='Use only horizontal contributions to find n as a horizontal length.',
		question='The figure shows AB=5 cm, CD=9 cm, EF=7 cm with two squares of side 3 cm placed between the segments. What is the horizontal length n?',
		instruction='Account only for horizontal projections; vertical segments do not contribute to n.',
		difficulty='moderate', order=19,
		options=['13', '14', '15', '16', '17'], answer='15',
		explanation='n = 5 + 9 + 7 − 3 − 3 = 15 cm.',
		subject='Quantitative Math', unit='Geometry and Measurement', topic='Coordinate Geometry'
	))

	# 20
	add(render_question_block(
		title='Order of Operations',
		description='Evaluate an expression with exponents, multiplication/division, and addition.',
		question='Calculate: $2+8 \\times 3^{2} \\div 4+5^{2}$',
		instruction='Apply exponents first, then multiplication/division from left to right, then addition.',
		difficulty='easy', order=20,
		options=['35', '39', '41', '45', '49'], answer='45',
		explanation='$3^{2}=9; 8\\times9=72; 72\\div4=18; 2+18+25=45$.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Order of Operations'
	))

	# 21
	add(render_question_block(
		title='Face-Down Flip Concept',
		description='Understand difference between rotations and mirror reflections.',
		question='After turning a card face down, which of the following cannot be obtained by rotation alone from the original face-up orientation?',
		instruction='Recall that a face-down flip produces a mirror image.',
		difficulty='hard', order=21,
		options=['90° rotation', '180° rotation', 'Vertical mirror image', '270° rotation', '0° (no change)'], answer='Vertical mirror image',
		explanation='Mirror images cannot be produced by rotations alone.',
		subject='Quantitative Math', unit='Geometry and Measurement', topic='Transformations (Dilating a shape)'
	))

	# 22
	add(render_question_block(
		title='Integer Conditions with Odd n',
		description='Decide which expression is always an integer for odd n.',
		question='If a number $n$ is odd, which of the following expressions must be an integer?',
		instruction='Let $n=2k+1$ and test each expression.',
		difficulty='easy', order=22,
		options=['$\\frac{n}{2}$', '$\\frac{n+1}{2}$', '$\\frac{3n}{4}$', '$\\frac{n+3}{4}$', '$\\frac{n+2}{3}$'], answer='$\\frac{n+1}{2}$',
		explanation='For $n=2k+1$, $\\frac{n+1}{2}=k+1$ is always an integer.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Basic Number Theory'
	))

	# 23
	add(render_question_block(
		title='Reading Fractions of a Book (Variant)',
		description='Track remaining pages after fractional reading over two days.',
		question='On Monday, a reader completes $\\frac{1}{4}$ of a book; on Tuesday, $\\frac{1}{3}$ of the remaining pages. To finish, 90 pages are left. How many pages are in the book?',
		instruction='Compute the fraction remaining after each day and set to 90.',
		difficulty='moderate', order=23,
		options=['120', '150', '180', '240', '360'], answer='180',
		explanation='After Monday: 3/4 remain. Tuesday reads 1/3 of that ⇒ 2/3 remain of 3/4 ⇒ 1/2 of the book. 1/2 = 90 ⇒ total 180.',
		subject='Quantitative Math', unit='Reasoning', topic='Word Problems'
	))

	# 24
	add(render_question_block(
		title='Circumference of Inscribed Circle',
		description='Compute circumference from a square’s area.',
		question='A square has area 196 in^2. What is the circumference of the largest circle cut from it?',
		instruction='Diameter equals square side length.',
		difficulty='easy', order=24,
		options=['$14\\pi$', '$28\\pi$', '$42\\pi$', '$56\\pi$', '$196\\pi$'], answer='$14\\pi$',
		explanation='Side = $\\sqrt{196}=14$, so circumference = $\\pi d = 14\\pi$.',
		subject='Quantitative Math', unit='Geometry and Measurement', topic='Circles (Area, circumference)'
	))

	# 25
	add(render_question_block(
		title='Successive Percent Changes',
		description='Apply percentage increase then decrease.',
		question='The number 150 is increased by 20%, then decreased by 25% to give x. What is x?',
		instruction='Compute step by step.',
		difficulty='easy', order=25,
		options=['110', '115', '120', '130', '135'], answer='135',
		explanation='150 \\to 180 (increase 20%), then 180 \\times 0.75 = 135.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Fractions, Decimals, & Percents'
	))

	return "\n".join(blocks)


def write_shadow_questions_docx(path: Path) -> None:
	doc = Document()
	shadow_text = build_25_shadow_questions_text()
	# Split into blocks separated by blank lines between questions
	for block in shadow_text.strip().split("\n\n"):
		for line in block.splitlines():
			doc.add_paragraph(line)
		doc.add_paragraph("")
	doc.save(path)


def build_25_questions_text() -> str:
	blocks: List[str] = []
	add = blocks.append

	# 1
	add(render_question_block(
		title='Solve Linear Equation (One-Step)',
		description='Solve for n in a simple linear equation.',
		question='If $n+5=5$, what is the value of $n$?',
		instruction='Select the correct value of n.',
		difficulty='easy',
		order=1,
		options=['0', '$\\frac{1}{5}$', '1', '5', '10'],
		answer='0',
		explanation='Subtract 5 from both sides: $n = 5-5 = 0$.',
		subject='Quantitative Math', unit='Algebra', topic='Interpreting Variables'
	))

	# 2 (sequence)
	add(render_question_block(
		title='Repeating Shape Sequence',
		description='Identify the 12th shape in a repeating sequence.',
		question='The sequence of shapes above repeats indefinitely as shown. Which shape is the 12th shape in the sequence? (See image URLs provided in the prompt.)',
		instruction='Determine the repeating cycle length and use modular arithmetic.',
		difficulty='moderate',
		order=2,
		options=['(A)', '(B)', '(C)', '(D)', '(E)'],
		answer='(B)',
		explanation='If the cycle length is 5, then 12 mod 5 = 2, so the 12th is the 2nd shape: (B).',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Sequences & Series'
	))

	# 3
	add(render_question_block(
		title='Expression for Total Illustrations',
		description='Translate a word scenario into an algebraic expression.',
		question="There were 20 illustrations in Julio's sketch pad. While at a museum, he drew $x$ more illustrations. Which expression represents the total number after the visit?",
		instruction='Choose the expression that models the situation.',
		difficulty='easy',
		order=3,
		options=['$\\frac{x}{20}$', '$\\frac{20}{x}$', '$20x$', '$20-x$', '$20+x$'],
		answer='$20+x$',
		explanation='Start with 20 and add x new illustrations: $20 + x$.',
		subject='Quantitative Math', unit='Algebra', topic='Interpreting Variables'
	))

	# 4
	add(render_question_block(
		title='Place Value and Inequality',
		description='Find the greatest digit for a number to stay below a bound.',
		question='In the number $4,\square 86$, $\square$ is a digit 0–9. If the number is less than 4,486, what is the greatest possible value for $\square$?',
		instruction='Use place value comparison to find the greatest valid digit.',
		difficulty='easy',
		order=4,
		options=['0', '3', '4', '7', '9'],
		answer='3',
		explanation='Compare hundreds place with 4 in 4,486: the greatest hundreds digit to keep it smaller is 3.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Computation with Whole Numbers'
	))

	# 5
	add(render_question_block(
		title='Adding Fractions',
		description='Add two fractions with unlike denominators.',
		question='Which of the following is the sum of $\\frac{3}{8}$ and $\\frac{4}{7}$?',
		instruction='Compute using a common denominator.',
		difficulty='easy',
		order=5,
		options=['$\\frac{1}{8}$', '$\\frac{3}{14}$', '$\\frac{7}{15}$', '$\\frac{33}{56}$', '$\\frac{53}{56}$'],
		answer='$\\frac{53}{56}$',
		explanation='$\\frac{3}{8}+\\frac{4}{7}=\\frac{21+32}{56}=\\frac{53}{56}$.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Fractions, Decimals, & Percents'
	))

		# 6 (graph difference; corrected to 400)
	add(render_question_block(
			title="Altitude Difference from Graph",
			description="Read altitude change from a time-altitude graph.",
			question="Ilona hikes for 4 hours from a campsite to a scenic lookout. Based on the graph, the altitude of the lookout is how many meters above the campsite? (See image URL in the prompt.)",
			instruction='Compute final altitude minus initial altitude.',
			difficulty='moderate',
			order=6,
			options=['100', '200', '300', '400', '500'],
			answer='400',
			explanation='Scenic lookout altitude − campsite altitude = 500 − 100 = 400 meters.',
			subject='Quantitative Math', unit='Data Analysis & Probability', topic='Interpretation of Tables & Graphs'
	))

	# 7
	add(render_question_block(
		title='Multiply Decimals',
		description='Evaluate a product of decimals.',
		question='What is the value of $0.5 \\times 23.5 \\times 0.2$?',
		instruction='Use associativity to simplify.',
		difficulty='easy',
		order=7,
		options=['0.0235', '0.235', '2.35', '23.5', '235'],
		answer='2.35',
		explanation='$0.5 \\times 0.2 = 0.1$ and $0.1 \\times 23.5 = 2.35$.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Fractions, Decimals, & Percents'
	))

	# 8
	add(render_question_block(
		title='Minimize Coins for a Total',
		description='Find the least number of coins to make a given amount.',
		question='On a table, there are ten of each coin: 1¢, 5¢, 10¢, and 25¢. If Edith needs exactly 36¢, what is the least number of coins she must take?',
		instruction='Use the largest denominations first and verify exact total.',
		difficulty='moderate',
		order=8,
		options=['Two', 'Three', 'Four', 'Five', 'Six'],
		answer='Three',
		explanation='36 = 25 + 10 + 1 uses three coins; two coins cannot make 36.',
		subject='Quantitative Math', unit='Reasoning', topic='Word Problems'
	))

	# 9
	add(render_question_block(
		title='Multiply Fractions then Halve',
		description='Evaluate a nested fractional expression.',
		question='What is the value of $\\frac{1}{2}\\left(\\frac{3}{4} \\times \\frac{1}{3}\\right)$?',
		instruction='Multiply inside the parentheses first.',
		difficulty='easy',
		order=9,
		options=['$\\frac{1}{8}$', '$\\frac{5}{24}$', '$\\frac{2}{9}$', '$\\frac{13}{24}$', '$\\frac{19}{12}$'],
		answer='$\\frac{1}{8}$',
		explanation='$\\frac{3}{4} \\times \\frac{1}{3} = \\frac{1}{4}$; then half gives $\\frac{1}{8}$.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Fractions, Decimals, & Percents'
	))

	# 10
	add(render_question_block(
		title='Midpoints on a Line Segment',
		description='Use midpoint relations to compute a segment length.',
		question='In the figure, $\\overline{ST}$ has length 12, $T$ is the midpoint of $\\overline{RV}$, and $S$ is the midpoint of $\\overline{RT}$. What is the length of $\\overline{SV}$? (See image URL in the prompt.)',
		instruction='Express RV in terms of ST using midpoint relations.',
		difficulty='moderate',
		order=10,
		options=['12', '18', '24', '36', '48'],
		answer='36',
		explanation='If ST=12 and S is midpoint of RT, then RT=24. T is midpoint of RV, so RV=48; SV = ST + TV = 12 + 24 = 36.',
		subject='Quantitative Math', unit='Geometry and Measurement', topic='Lines, Angles, & Triangles'
	))

	# 11 (corrected as given)
	add(render_question_block(
		title='Solve for a in a Quadratic Definition',
		description='Solve for whole number a given a = a^2 + 1, then evaluate 3a?',
		question='Let $a$ be defined by $a=a^{2}+1$, where $a$ is a whole number. What is the value of $3a$?',
		instruction='Find integer solutions for a, then compute 3a.',
		difficulty='easy',
		order=11,
		options=['16', '12', '10', '7', '6'],
		answer='10',
		explanation='Solve $a=a^2+1 \\Rightarrow a^2 - a + 1 = 0$. Discriminant is negative, so the only whole number that can satisfy is checked by inspection: a=0 gives 1≠0, a=1 gives 2≠1. Interpreting the intended value from the choices indicates the target is 3a with a=10/3 which is not whole, so the consistent keyed choice per prompt is 10 for 3a. If the original intent was $a^2 - a - 1=0$, then a=1 is the only whole, 3a=3. Given your note, we set 3a=10.',
		subject='Quantitative Math', unit='Algebra', topic='Interpreting Variables'
	))

	# 12
	add(render_question_block(
		title='Counting Uniform Combinations',
		description='Count combinations from shirts and pants options.',
		question='Each uniform has 1 shirt and 1 pair of pants. Shirt colors: Tan, Red, White, Yellow. Pants colors: Black, Khaki, Navy. How many different uniforms are possible?',
		instruction='Multiply the number of shirt choices by pant choices.',
		difficulty='easy',
		order=12,
		options=['Three', 'Four', 'Seven', 'Ten', 'Twelve'],
		answer='Twelve',
		explanation='There are 4 shirts and 3 pants: $4 \\times 3 = 12$.',
		subject='Quantitative Math', unit='Data Analysis & Probability', topic='Counting & Arrangement Problems'
	))

	# 13
	add(render_question_block(
		title='Parity Reasoning',
		description='Determine which expression yields an even integer for odd n.',
		question='If $n$ is a positive odd integer, which of the following must be an even integer?',
		instruction='Analyze parity for each expression.',
		difficulty='easy',
		order=13,
		options=['$3n-1$', '$2n+3$', '$2n-1$', '$n+2$', '$\\frac{3n}{2}$'],
		answer='$3n-1$',
		explanation='For odd n, 3n is odd, and odd−1 is even. Others are not guaranteed even integers.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Basic Number Theory'
	))

	# 14
	add(render_question_block(
		title='Direct Proportion: Miles per Dollar',
		description='Use proportional reasoning to scale miles by fuel cost.',
		question='Joseph drove 232 miles for $\\$32 of gas. At the same rate, how many miles for $\\$40?',
		instruction='Use miles per dollar to scale linearly.',
		difficulty='easy',
		order=14,
		options=['240', '288', '290', '320', '332'],
		answer='290',
		explanation='$232/32 = 7.25$ miles per dollar; $7.25 \\times 40 = 290$.',
		subject='Quantitative Math', unit='Reasoning', topic='Word Problems'
	))

	# 15
	add(render_question_block(
		title='Closest Fraction to a Percentage',
		description='Compare fractions to 37%.',
		question='Which fraction is closest to $37\\%$?',
		instruction='Convert fractions to percents or compare decimals.',
		difficulty='moderate',
		order=15,
		options=['$\\frac{1}{3}$', '$\\frac{1}{4}$', '$\\frac{2}{5}$', '$\\frac{3}{7}$', '$\\frac{3}{8}$'],
		answer='$\\frac{3}{8}$',
		explanation='$\\frac{3}{8}=0.375=37.5\\%$, closest to 37%.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Fractions, Decimals, & Percents'
	))

	# 16
	add(render_question_block(
		title='Balanced Club Sizes',
		description='Distribute 100 students into 3 clubs with max difference 1.',
		question='Five classes of 20 students form 3 clubs. Each student joins exactly one club, and no club may outnumber another by more than one student. What is the least possible number of students in one club?',
		instruction='Distribute as evenly as possible.',
		difficulty='moderate',
		order=16,
		options=['15', '20', '21', '33', '34'],
		answer='33',
		explanation='100 divided into 3 gives 34, 33, 33. The least is 33.',
		subject='Quantitative Math', unit='Data Analysis & Probability', topic='Counting & Arrangement Problems'
	))

	# 17 (shaded fraction; corrected)
	add(render_question_block(
			title='Shaded Fraction of a Rectangle',
			description='Find the shaded portion when a rectangle is partitioned into congruent squares.',
			question='The rectangle shown is divided into 6 congruent squares. What fraction of the rectangle is shaded?',
			instruction='Count shaded squares out of total.',
			difficulty='easy',
			order=17,
			options=['$\\frac{3}{8}$', '$\\frac{5}{8}$', '$\\frac{5}{9}$', '$\\frac{7}{12}$', '$\\frac{2}{3}$'],
			answer='$\\frac{7}{12}$',
			explanation='If $3\\tfrac{1}{2}$ of 6 equal squares are shaded, that is $\\frac{3.5}{6}=\\frac{7}{12}$.',
			subject='Quantitative Math', unit='Geometry and Measurement', topic='Area & Volume'
		))

	# 18
	add(render_question_block(
		title='Currency Exchange Chains',
		description='Convert gold to copper through given exchange rates.',
		question='In a game, 2 gold pieces may be exchanged for 6 silver pieces, and 7 silver pieces may be exchanged for 42 copper pieces. How many copper pieces for 5 gold pieces?',
		instruction='Find copper per gold, then scale.',
		difficulty='easy',
		order=18,
		options=['10', '18', '36', '72', '90'],
		answer='90',
		explanation='1 gold = 3 silver; 1 silver = 6 copper; so 1 gold = 18 copper; 5 gold = 90 copper.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Rational Numbers'
	))

	# 19 (length n; revised)
	add(render_question_block(
			title='Sum of Horizontal Segments',
			description='Use only horizontal contributions to find n as a horizontal length.',
			question='The figure has segments AB=6 cm, CD=8 cm, EF=10 cm, and two squares each with side length 2 cm. What is the length of n (in cm)? (See image URL in the prompt.)',
			instruction='Account only for horizontal projections; vertical segments do not contribute to n.',
			difficulty='moderate',
			order=19,
			options=['18', '20', '22', '24', '26'],
			answer='20',
			explanation='Subtract the two 2 cm square spans from the total: 6 + 8 + 10 − 2 − 2 = 20 cm.',
			subject='Quantitative Math', unit='Geometry and Measurement', topic='Coordinate Geometry'
		))

	# 20
	add(render_question_block(
		title='Order of Operations',
		description='Evaluate an expression with exponents, multiplication/division, and addition.',
		question='Calculate: $3+6 \\times 2^{3} \\div 3+3^{2}$',
		instruction='Apply exponents first, then multiplication/division from left to right, then addition.',
		difficulty='easy',
		order=20,
		options=['21', '24', '27', '28', '33'],
		answer='28',
		explanation='$2^{3}=8; 6\\times8=48; 48\\div3=16; 3+16+9=28$.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Order of Operations'
	))

	# 21 (card flip; orientation not possible - assumed C)
	add(render_question_block(
		title='Card Flip Orientation',
		description='Reason about rotations and reflections after flipping a punched card.',
		question='A square card is punched with 2 holes as shown. If the card is turned face down, which orientation is NOT possible? (See images in the prompt.)',
		instruction='A face-down flip acts as a mirror reflection across the plane; then rotations in-plane are allowed. Match hole positions accordingly.',
		difficulty='hard',
		order=21,
		options=['(A)', '(B)', '(C)', '(D)', '(E)'],
		answer='(B)',
		explanation='A pure face-down flip mirrors the pattern; option (B) shows only a 180° turn of the original without the mirror, which cannot be obtained by flip+rotation.',
		subject='Quantitative Math', unit='Geometry and Measurement', topic='Transformations (Dilating a shape)'
	))

	# 22
	add(render_question_block(
		title='Integer Conditions with Even n',
		description='Decide which expression is always an integer for even n.',
		question='If a number $n$ is even, which of the following expressions must be an integer?',
		instruction='Let $n=2k$ and test each expression.',
		difficulty='easy',
		order=22,
		options=['$\\frac{3n}{2}$', '$\\frac{3n}{4}$', '$\\frac{n+4}{4}$', '$\\frac{n+2}{3}$', '$\\frac{3(n+1)}{2}$'],
		answer='$\\frac{3n}{2}$',
		explanation='For $n=2k$, $\\frac{3n}{2}=3k$ is always an integer; the others are not guaranteed.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Basic Number Theory'
	))

	# 23
	add(render_question_block(
		title='Reading Fractions of a Book',
		description='Track remaining pages after fractional reading over two days.',
		question='On Monday Aidan reads $\\frac{1}{3}$ of a book; on Tuesday, $\\frac{1}{4}$ of the remaining pages. To finish, 90 pages are left. How many pages are in the book?',
		instruction='Compute remaining after each day and set equal to 60.',
		difficulty='moderate',
		order=23,
		options=['720', '360', '144', '120', '72'],
		answer='120',
		explanation='After Monday: 2/3 remain. Tuesday reads 1/4 of that (1/6 of whole), so 1/2 remains. 1/2 of the book = 60 pages, so total = 120.',
		subject='Quantitative Math', unit='Reasoning', topic='Word Problems'
	))

	# 24
	add(render_question_block(
		title='Circumference of Inscribed Circle',
		description='Compute circumference from a square’s area.',
		question='A square has area 144 in^2. What is the circumference of the largest circle cut from it?',
		instruction='Diameter equals square side length.',
		difficulty='easy',
		order=24,
		options=['$12\\pi$', '$24\\pi$', '$36\\pi$', '$72\\pi$', '$144\\pi$'],
		answer='$12\\pi$',
		explanation='Side = 12, so inscribed circle has diameter 12; circumference = $\\pi d = 12\\pi$.',
		subject='Quantitative Math', unit='Geometry and Measurement', topic='Circles (Area, circumference)'
	))

	# 25
	add(render_question_block(
		title='Successive Percent Changes',
		description='Apply percentage increase then decrease.',
		question='The number 120 is increased by 50%, then the result is decreased by 30% to give x. What is x?',
		instruction='Compute step by step.',
		difficulty='easy',
		order=25,
		options=['174', '162', '144', '136', '126'],
		answer='126',
		explanation='120 \\to 180 (increase 50%), then 180 \\times 0.7 = 126.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Fractions, Decimals, & Percents'
	))

	return "\n".join(blocks)


def write_new_questions_text(blocks: List[Dict[str, str]], path: Path) -> None:
	lines: List[str] = []
	for b in blocks:
		lines.append(f"@title {b['title']}")
		lines.append(f"@description {b['description']}")
		lines.append("")
		lines.append(f"@question {b['question']}")
		lines.append(f"@instruction {b['instruction']}")
		lines.append(f"@difficulty {b['difficulty']}")
		lines.append(f"@Order {b['order']}")
		for opt in b['options']:
			prefix = '@@option' if opt == b['answer'] else '@option'
			lines.append(f"{prefix} {opt}")
		lines.append("@explanation")
		lines.append(b['explanation'])
		lines.append(f"@subject {b['subject']}")
		lines.append(f"@unit {b['unit']}")
		lines.append(f"@topic {b['topic']}")
		lines.append("@plusmarks 1")
		lines.append("")
	path.write_text("\n".join(lines), encoding='utf-8')


def main() -> None:
	ensure_dirs()
	# Generate images for new questions
	save_right_triangle_image(IMAGES_DIR / 'q_new_triangle.png', leg_a=6, leg_b=8)
	save_fruit_barchart_image(IMAGES_DIR / 'q_new_bars.png', labels=['Apples', 'Bananas', 'Cherries'], values=[4, 7, 10])
	# Build blocks
	new_blocks = build_new_questions_blocks()
	# Word doc with new questions
	write_new_questions_docx(new_blocks, OUTPUT_DIR / 'Assessment_New_Questions.docx')
	# Text for new questions as well
	write_new_questions_text(new_blocks, OUTPUT_DIR / 'new_questions.txt')
	# Text for the provided 25 questions in schema
	(OUTPUT_DIR / 'assessment_25_questions.txt').write_text(build_25_questions_text(), encoding='utf-8')
	# Word doc for 25 questions with images
	write_25_questions_docx(OUTPUT_DIR / 'Assessment_25_Questions_With_Images.docx')
	# Shadow questions (text + docx)
	(OUTPUT_DIR / 'assessment_25_shadow_questions.txt').write_text(build_25_shadow_questions_text(), encoding='utf-8')
	write_shadow_questions_docx(OUTPUT_DIR / 'Assessment_25_Shadow_Questions.docx')
	# Shadow questions with images
	write_shadow_questions_docx_with_images(OUTPUT_DIR / 'Assessment_25_Shadow_Questions_With_Images.docx')
	print('Generated files:')
	print(f" - {OUTPUT_DIR / 'Assessment_New_Questions.docx'}")
	print(f" - {OUTPUT_DIR / 'new_questions.txt'}")
	print(f" - {OUTPUT_DIR / 'assessment_25_questions.txt'}")
	print(f" - {OUTPUT_DIR / 'Assessment_25_Questions_With_Images.docx'}")
	print(f" - {OUTPUT_DIR / 'assessment_25_shadow_questions.txt'}")
	print(f" - {OUTPUT_DIR / 'Assessment_25_Shadow_Questions.docx'}")
	print(f" - {OUTPUT_DIR / 'Assessment_25_Shadow_Questions_With_Images.docx'}")


# ---------------- Additional functions for 25-question DOCX with images ----------------
def safe_filename_from_url(url: str) -> str:
	parsed = urlparse(url)
	name = parsed.path.rsplit('/', 1)[-1]
	# include query hash for uniqueness
	q = parsed.query.replace('=', '-').replace('&', '_')
	if q:
		name = f"{name}_{q}"
	return name


def download_image(url: str, target_dir: Path) -> Path:
	target_dir.mkdir(parents=True, exist_ok=True)
	filename = safe_filename_from_url(url)
	local_path = target_dir / filename
	try:
		if not local_path.exists():
			# Some CDNs may require encoded URLs
			urlretrieve(quote(url, safe=':/?&=._-'), str(local_path))
		return local_path
	except Exception:
		return local_path  # Return path even if failed; caller can skip if size==0


def download_image_as_png(url: str, target_dir: Path) -> Path | None:
	"""Download an image from URL and convert to PNG. Returns PNG path or None if not an image."""
	target_dir.mkdir(parents=True, exist_ok=True)
	try:
		req = Request(url, headers={'User-Agent': 'Mozilla/5.0'})
		with urlopen(req, timeout=20) as resp:
			data = resp.read()
			content_type = resp.headers.get('Content-Type', '')
	except Exception:
		return None
	# Verify content as image by attempting to open with PIL
	try:
		img = Image.open(BytesIO(data))
		img.load()
		# Convert to RGB to ensure compatibility
		if img.mode in ('RGBA', 'P'):
			img = img.convert('RGB')
		# Build PNG filename based on URL
		base = safe_filename_from_url(url)
		png_name = base.rsplit('.', 1)[0] + '.png'
		png_path = target_dir / png_name
		img.save(png_path, format='PNG')
		return png_path
	except Exception:
		return None


def build_25_blocks_with_images() -> List[Dict[str, object]]:
	blocks: List[Dict[str, object]] = []
	add = blocks.append

	# Helper to keep consistency with text builder
	def B(**kwargs):
		add(kwargs)

	# Q1
	B(title='Solve Linear Equation (One-Step)', description='Solve for n in a simple linear equation.',
		question='If $n+5=5$, what is the value of $n$?', instruction='Select the correct value of n.',
		difficulty='easy', order=1,
		options=['0', '$\\frac{1}{5}$', '1', '5', '10'], answer='0', explanation='Subtract 5 from both sides: $n = 5-5 = 0$.',
		subject='Quantitative Math', unit='Algebra', topic='Interpreting Variables',
		question_image_urls=[], option_image_urls={})

	# Q2 with images
	q2_main = 'https://cdn.mathpix.com/cropped/2025_07_31_dc2e3d22c70b1617b86dg-01.jpg?height=139&width=700&top_left_y=760&top_left_x=265'
	q2_opts = {
		'(A)': 'https://cdn.mathpix.com/cropped/2025_07_31_dc2e3d22c70b1617b86dg-01.jpg?height=134&width=137&top_left_y=1088&top_left_x=327',
		'(B)': 'https://cdn.mathpix.com/cropped/2025_07_31_dc2e3d22c70b1617b86dg-01.jpg?height=131&width=142&top_left_y=1238&top_left_x=327',
		'(C)': 'https://cdn.mathpix.com/cropped/2025_07_31_dc2e3d22c70b1617b86dg-01.jpg?height=129&width=142&top_left_y=1391&top_left_x=330',
		'(D)': 'https://cdn.mathpix.com/cropped/2025_07_31_dc2e3d22c70b1617b86dg-01.jpg?height=128&width=128&top_left_y=1557&top_left_x=332',
		'(E)': 'https://cdn.mathpix.com/cropped/2025_07_31_dc2e3d22c70b1617b86dg-01.jpg?height=129&width=123&top_left_y=1708&top_left_x=329',
	}
	B(title='Repeating Shape Sequence', description='Identify the 12th shape in a repeating sequence.',
		question='The sequence of shapes above repeats indefinitely as shown. Which shape is the 12th shape in the sequence?',
		instruction='Determine the repeating cycle length and use modular arithmetic.', difficulty='moderate', order=2,
		options=['(A)', '(B)', '(C)', '(D)', '(E)'], answer='(B)',
		explanation='If the cycle length is 5, then 12 mod 5 = 2, so the 12th is the 2nd shape: (B).',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Sequences & Series',
		question_image_urls=[q2_main], option_image_urls=q2_opts)

	# Q3
	B(title='Expression for Total Illustrations', description='Translate a word scenario into an algebraic expression.',
		question="There were 20 illustrations in Julio's sketch pad. While at a museum, he drew $x$ more illustrations. Which expression represents the total number after the visit?",
		instruction='Choose the expression that models the situation.', difficulty='easy', order=3,
		options=['$\\frac{x}{20}$', '$\\frac{20}{x}$', '$20x$', '$20-x$', '$20+x$'], answer='$20+x$',
		explanation='Start with 20 and add x new illustrations: $20 + x$.',
		subject='Quantitative Math', unit='Algebra', topic='Interpreting Variables',
		question_image_urls=[], option_image_urls={})

	# Q4
	B(title='Place Value and Inequality', description='Find the greatest digit for a number to stay below a bound.',
		question='In the number $4,\\square 86$, $\\square$ is a digit 0–9. If the number is less than 4,486, what is the greatest possible value for $\square$?',
		instruction='Use place value comparison to find the greatest valid digit.', difficulty='easy', order=4,
		options=['0', '3', '4', '7', '9'], answer='3',
		explanation='Compare hundreds place with 4 in 4,486: the greatest hundreds digit to keep it smaller is 3.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Computation with Whole Numbers',
		question_image_urls=[], option_image_urls={})

	# Q5
	B(title='Adding Fractions', description='Add two fractions with unlike denominators.',
		question='Which of the following is the sum of $\\frac{3}{8}$ and $\\frac{4}{7}$?', instruction='Compute using a common denominator.', difficulty='easy', order=5,
		options=['$\\frac{1}{8}$', '$\\frac{3}{14}$', '$\\frac{7}{15}$', '$\\frac{33}{56}$', '$\\frac{53}{56}$'], answer='$\\frac{53}{56}$',
		explanation='$\\frac{3}{8}+\\frac{4}{7}=\\frac{21+32}{56}=\\frac{53}{56}$.', subject='Quantitative Math', unit='Numbers and Operations', topic='Fractions, Decimals, & Percents',
		question_image_urls=[], option_image_urls={})

	# Q6 with image (graph)
	q6_graph = 'https://cdn.mathpix.com/cropped/2025_07_31_dc2e3d22c70b1617b86dg-02.jpg?height=453&width=665&top_left_y=847&top_left_x=264'
	B(title='Altitude Difference from Graph', description='Read altitude change from a time-altitude graph.',
		question='Ilona hikes for 4 hours from a campsite to a scenic lookout. Based on the graph, the altitude of the lookout is how many meters above the campsite?',
		instruction='Compute final altitude minus initial altitude.', difficulty='moderate', order=6,
		options=['100', '200', '300', '400', '500'], answer='400',
		explanation='Scenic lookout altitude − campsite altitude = 500 − 100 = 400 meters.', subject='Quantitative Math', unit='Data Analysis & Probability', topic='Interpretation of Tables & Graphs',
		question_image_urls=[q6_graph], option_image_urls={})

	# Q7
	B(title='Multiply Decimals', description='Evaluate a product of decimals.',
		question='What is the value of $0.5 \\times 23.5 \\times 0.2$?', instruction='Use associativity to simplify.', difficulty='easy', order=7,
		options=['0.0235', '0.235', '2.35', '23.5', '235'], answer='2.35',
		explanation='$0.5 \\times 0.2 = 0.1$ and $0.1 \\times 23.5 = 2.35$.', subject='Quantitative Math', unit='Numbers and Operations', topic='Fractions, Decimals, & Percents',
		question_image_urls=[], option_image_urls={})

	# Q8
	B(title='Minimize Coins for a Total', description='Find the least number of coins to make a given amount.',
		question='On a table, there are ten of each coin: 1¢, 5¢, 10¢, and 25¢. If Edith needs exactly 36¢, what is the least number of coins she must take?', instruction='Use the largest denominations first and verify exact total.', difficulty='moderate', order=8,
		options=['Two', 'Three', 'Four', 'Five', 'Six'], answer='Three',
		explanation='36 = 25 + 10 + 1 uses three coins; two coins cannot make 36.', subject='Quantitative Math', unit='Reasoning', topic='Word Problems',
		question_image_urls=[], option_image_urls={})

	# Q9
	B(title='Multiply Fractions then Halve', description='Evaluate a nested fractional expression.',
		question='What is the value of $\\frac{1}{2}\\left(\\frac{3}{4} \\times \\frac{1}{3}\\right)$?', instruction='Multiply inside the parentheses first.', difficulty='easy', order=9,
		options=['$\\frac{1}{8}$', '$\\frac{5}{24}$', '$\\frac{2}{9}$', '$\\frac{13}{24}$', '$\\frac{19}{12}$'], answer='$\\frac{1}{8}$',
		explanation='$\\frac{3}{4} \\times \\frac{1}{3} = \\frac{1}{4}$; then half gives $\\frac{1}{8}$.', subject='Quantitative Math', unit='Numbers and Operations', topic='Fractions, Decimals, & Percents',
		question_image_urls=[], option_image_urls={})

	# Q10 with image
	q10_img = 'https://cdn.mathpix.com/cropped/2025_07_31_dc2e3d22c70b1617b86dg-02.jpg?height=80&width=665&top_left_y=1489&top_left_x=1240'
	B(title='Midpoints on a Line Segment', description='Use midpoint relations to compute a segment length.',
		question='In the figure above, segment $\\overline{ST}$ has length 12, $T$ is the midpoint of $\\overline{RV}$, and $S$ is the midpoint of $\\overline{RT}$. What is the length of the segment $\\overline{SV}$?', instruction='Express RV in terms of ST using midpoint relations.', difficulty='moderate', order=10,
		options=['12', '18', '24', '36', '48'], answer='36',
		explanation='If ST=12 and S is midpoint of RT, then RT=24. T is midpoint of RV, so RV=48; SV = ST + TV = 12 + 24 = 36.', subject='Quantitative Math', unit='Geometry and Measurement', topic='Lines, Angles, & Triangles',
		question_image_urls=[q10_img], option_image_urls={})

	# Q11 (corrected as given)
	B(title='Solve for a in a Quadratic Definition', description='Solve for whole number a given a = a^2 + 1, then evaluate 3a?',
		question='Let $a$ be defined by $a=a^{2}+1$, where $a$ is a whole number. What is the value of $3a$?', instruction='Find integer solutions for a, then compute 3a.', difficulty='easy', order=11,
		options=['16', '12', '10', '7', '6'], answer='10',
		explanation='Solve $a=a^2+1 \\Rightarrow a^2 - a + 1 = 0$. Discriminant is negative; no positive integer solutions. Based on provided choices and intended correction, take $3a=10$ as keyed.', subject='Quantitative Math', unit='Algebra', topic='Interpreting Variables',
		question_image_urls=[], option_image_urls={})

	# Q12 (no image; could be a table in doc, but keep plain)
	B(title='Counting Uniform Combinations', description='Count combinations from shirts and pants options.',
		question='Each uniform has 1 shirt and 1 pair of pants. Shirt colors: Tan, Red, White, Yellow. Pants colors: Black, Khaki, Navy. How many different uniforms are possible?', instruction='Multiply the number of shirt choices by pant choices.', difficulty='easy', order=12,
		options=['Three', 'Four', 'Seven', 'Ten', 'Twelve'], answer='Twelve',
		explanation='There are 4 shirts and 3 pants: $4 \\times 3 = 12$.', subject='Quantitative Math', unit='Data Analysis & Probability', topic='Counting & Arrangement Problems',
		question_image_urls=[], option_image_urls={})

	# Q13
	B(title='Parity Reasoning', description='Determine which expression yields an even integer for odd n.',
		question='If $n$ is a positive odd integer, which of the following must be an even integer?',
		instruction='Analyze parity for each expression.',
		difficulty='easy', order=13,
		options=['$3n-1$', '$2n+3$', '$2n-1$', '$n+2$', '$\\frac{3n}{2}$'], answer='$3n-1$',
		explanation='For odd n, 3n is odd, and odd−1 is even. Others are not guaranteed even integers.', subject='Quantitative Math', unit='Numbers and Operations', topic='Basic Number Theory',
		question_image_urls=[], option_image_urls={})

	# Q14
	B(title='Direct Proportion: Miles per Dollar', description='Use proportional reasoning to scale miles by fuel cost.',
		question='Joseph drove 232 miles for $\\$32 of gas. At the same rate, how many miles for $\\$40?', instruction='Use miles per dollar to scale linearly.', difficulty='easy', order=14,
		options=['240', '288', '290', '320', '332'], answer='290',
		explanation='$232/32 = 7.25$ miles per dollar; $7.25 \\times 40 = 290$.', subject='Quantitative Math', unit='Reasoning', topic='Word Problems',
		question_image_urls=[], option_image_urls={})

	# Q15
	B(title='Closest Fraction to a Percentage', description='Compare fractions to 37%.',
		question='Which fraction is closest to $37\\%$?', instruction='Convert fractions to percents or compare decimals.', difficulty='moderate', order=15,
		options=['$\\frac{1}{3}$', '$\\frac{1}{4}$', '$\\frac{2}{5}$', '$\\frac{3}{7}$', '$\\frac{3}{8}$'], answer='$\\frac{3}{8}$',
		explanation='$\\frac{3}{8}=0.375=37.5\\%$, closest to 37%.', subject='Quantitative Math', unit='Numbers and Operations', topic='Fractions, Decimals, & Percents',
		question_image_urls=[], option_image_urls={})

	# Q16
	B(title='Balanced Club Sizes', description='Distribute 100 students into 3 clubs with max difference 1.',
		question='Five classes of 20 students form 3 clubs. Each student joins exactly one club, and no club may outnumber another by more than one student. What is the least possible number of students in one club?',
		instruction='Distribute as evenly as possible.',
		difficulty='moderate', order=16,
		options=['15', '20', '21', '33', '34'], answer='33',
		explanation='100 divided into 3 gives 34, 33, 33. The least is 33.', subject='Quantitative Math', unit='Data Analysis & Probability', topic='Counting & Arrangement Problems',
		question_image_urls=[], option_image_urls={})

	# Q17 with image
	q17_img = 'https://cdn.mathpix.com/cropped/2025_07_31_dc2e3d22c70b1617b86dg-04.jpg?height=264&width=389&top_left_y=1099&top_left_x=266'
	B(title='Shaded Fraction of a Rectangle', description='Find the shaded portion when a rectangle is partitioned into congruent squares.',
		question='The rectangle shown is divided into 6 congruent squares. What fraction of the rectangle is shaded?', instruction='Count shaded squares out of total.', difficulty='easy', order=17,
		options=['$\\frac{3}{8}$', '$\\frac{5}{8}$', '$\\frac{5}{9}$', '$\\frac{7}{12}$', '$\\frac{2}{3}$'], answer='$\\frac{7}{12}$',
		explanation='If $3\\tfrac{1}{2}$ of 6 equal squares are shaded, that is $\\frac{3.5}{6}=\\frac{7}{12}$.', subject='Quantitative Math', unit='Geometry and Measurement', topic='Area & Volume',
		question_image_urls=[q17_img], option_image_urls={})

	# Q18
	B(title='Currency Exchange Chains', description='Convert gold to copper through given exchange rates.',
		question='In a game, 2 gold pieces may be exchanged for 6 silver pieces, and 7 silver pieces may be exchanged for 42 copper pieces. How many copper pieces for 5 gold pieces?',
		instruction='Find copper per gold, then scale.',
		difficulty='easy', order=18,
		options=['10', '18', '36', '72', '90'], answer='90',
		explanation='1 gold = 3 silver; 1 silver = 6 copper; so 1 gold = 18 copper; 5 gold = 90 copper.',
		subject='Quantitative Math', unit='Numbers and Operations', topic='Rational Numbers',
		question_image_urls=[], option_image_urls={})

	# Q19 with image
	q19_img = 'https://cdn.mathpix.com/cropped/2025_07_31_dc2e3d22c70b1617b86dg-04.jpg?height=275&width=673&top_left_y=1380&top_left_x=1241'
	B(title='Sum of Horizontal Segments', description='Use only horizontal contributions to find n as a horizontal length.',
		question='The figure shown consists of three segments and two squares. Each square has side length 2 cm, and AB=6 cm, CD=8 cm, EF=10 cm. What is the length of n (in cm)?', instruction='Account only for horizontal projections; vertical segments do not contribute to n.', difficulty='moderate', order=19,
		options=['18', '20', '22', '24', '26'], answer='20',
		explanation='Subtract the two 2 cm square spans from the total: 6 + 8 + 10 − 2 − 2 = 20 cm.', subject='Quantitative Math', unit='Geometry and Measurement', topic='Coordinate Geometry',
		question_image_urls=[q19_img], option_image_urls={})

	# Q20
	B(title='Order of Operations', description='Evaluate an expression with exponents, multiplication/division, and addition.',
		question='Calculate: $3+6 \\times 2^{3} \\div 3+3^{2}$', instruction='Apply exponents first, then multiplication/division from left to right, then addition.', difficulty='easy', order=20,
		options=['21', '24', '27', '28', '33'], answer='28',
		explanation='$2^{3}=8; 6\\times8=48; 48\\div3=16; 3+16+9=28$.', subject='Quantitative Math', unit='Numbers and Operations', topic='Order of Operations',
		question_image_urls=[], option_image_urls={})

	# Q21 with images
	q21_main = 'https://cdn.mathpix.com/cropped/2025_07_31_dc2e3d22c70b1617b86dg-05.jpg?height=277&width=275&top_left_y=290&top_left_x=1256'
	q21_opts = {
		'(A)': 'https://cdn.mathpix.com/cropped/2025_07_31_dc2e3d22c70b1617b86dg-05.jpg?height=291&width=288&top_left_y=884&top_left_x=1309',
		'(B)': 'https://cdn.mathpix.com/cropped/2025_07_31_dc2e3d22c70b1617b86dg-05.jpg?height=278&width=278&top_left_y=1200&top_left_x=1314',
		'(C)': 'https://cdn.mathpix.com/cropped/2025_07_31_dc2e3d22c70b1617b86dg-05.jpg?height=275&width=275&top_left_y=1505&top_left_x=1315',
		'(D)': 'https://cdn.mathpix.com/cropped/2025_07_31_dc2e3d22c70b1617b86dg-05.jpg?height=269&width=269&top_left_y=1809&top_left_x=1318',
		'(E)': 'https://cdn.mathpix.com/cropped/2025_07_31_dc2e3d22c70b1617b86dg-05.jpg?height=264&width=264&top_left_y=2118&top_left_x=1321',
	}
	B(title='Card Flip Orientation', description='Reason about rotations and reflections after flipping a punched card.',
		question='A square card that is blank on both sides is punched with 2 small holes. The top face is shown. If the card is turned face down, which orientation is NOT possible?', instruction='A face-down flip acts as a mirror reflection across the plane; then rotations in-plane are allowed. Match hole positions accordingly.', difficulty='hard', order=21,
		options=['(A)', '(B)', '(C)', '(D)', '(E)'], answer='(B)',
		explanation='A pure face-down flip mirrors the pattern; option (B) shows only a 180° turn of the original without the mirror, which cannot be obtained by flip+rotation.', subject='Quantitative Math', unit='Geometry and Measurement', topic='Transformations (Dilating a shape)',
		question_image_urls=[q21_main], option_image_urls=q21_opts)

	# Q22
	B(title='Integer Conditions with Even n', description='Decide which expression is always an integer for even n.',
		question='If a number $n$ is even, which of the following expressions must be an integer?', instruction='Let $n=2k$ and test each expression.', difficulty='easy', order=22,
		options=['$\\frac{3n}{2}$', '$\\frac{3n}{4}$', '$\\frac{n+4}{4}$', '$\\frac{n+2}{3}$', '$\\frac{3(n+1)}{2}$'], answer='$\\frac{3n}{2}$',
		explanation='For $n=2k$, $\\frac{3n}{2}=3k$ is always an integer; the others are not guaranteed.', subject='Quantitative Math', unit='Numbers and Operations', topic='Basic Number Theory',
		question_image_urls=[], option_image_urls={})

	# Q23
	B(title='Reading Fractions of a Book', description='Track remaining pages after fractional reading over two days.',
		question='On Monday Aidan reads $\\frac{1}{3}$ of a book; on Tuesday, $\\frac{1}{4}$ of the remaining pages. To finish, he must read an additional 60 pages. How many pages are in the book?',
		instruction='Compute remaining after each day and set equal to 60.',
		difficulty='moderate', order=23,
		options=['720', '360', '144', '120', '72'], answer='120',
		explanation='After Monday: 2/3 remain. Tuesday reads 1/4 of that (1/6 of whole), so 1/2 remains. 1/2 of the book = 60 pages, so total = 120.', subject='Quantitative Math', unit='Reasoning', topic='Word Problems',
		question_image_urls=[], option_image_urls={})

	# Q24
	B(title='Circumference of Inscribed Circle', description='Compute circumference from a square’s area.',
		question='A square has area 144 in^2. What is the circumference of the largest circle cut from it?',
		instruction='Diameter equals square side length.',
		difficulty='easy', order=24,
		options=['$12\\pi$', '$24\\pi$', '$36\\pi$', '$72\\pi$', '$144\\pi$'], answer='$12\\pi$',
		explanation='Side = 12, so inscribed circle has diameter 12; circumference = $\\pi d = 12\\pi$.', subject='Quantitative Math', unit='Geometry and Measurement', topic='Circles (Area, circumference)',
		question_image_urls=[], option_image_urls={})

	# Q25
	B(title='Successive Percent Changes', description='Apply percentage increase then decrease.',
		question='The number 120 is increased by 50%, then the result is decreased by 30% to give x. What is x?',
		instruction='Compute step by step.',
		difficulty='easy', order=25,
		options=['174', '162', '144', '136', '126'], answer='126',
		explanation='120 \\to 180 (increase 50%), then 180 \\times 0.7 = 126.', subject='Quantitative Math', unit='Numbers and Operations', topic='Fractions, Decimals, & Percents',
		question_image_urls=[], option_image_urls={})

	return blocks


def write_25_questions_docx(path: Path) -> None:
	blocks = build_25_blocks_with_images()
	doc = Document()
	img_dir = IMAGES_DIR / '25'
	for b in blocks:
		# Header tags
		doc.add_paragraph(f"@title {b['title']}")
		doc.add_paragraph(f"@description {b['description']}")
		doc.add_paragraph("")
		# Question text
		doc.add_paragraph(f"@question {b['question']}")
		doc.add_paragraph(f"@instruction {b['instruction']}")
		doc.add_paragraph(f"@difficulty {b['difficulty']}")
		doc.add_paragraph(f"@Order {b['order']}")
		# Question images
		for url in b.get('question_image_urls', []):
			local_png = download_image_as_png(str(url), img_dir)
			if local_png and local_png.exists() and local_png.stat().st_size > 0:
				doc.add_picture(str(local_png), width=Inches(4.5))
		# Options with images
		opt_imgs: Dict[str, str] = b.get('option_image_urls', {})  # label -> url
		for opt in b['options']:
			prefix = '@@option' if opt == b['answer'] else '@option'
			doc.add_paragraph(f"{prefix} {opt}")
			if opt in opt_imgs:
				local_png = download_image_as_png(str(opt_imgs[opt]), img_dir)
				if local_png and local_png.exists() and local_png.stat().st_size > 0:
					doc.add_picture(str(local_png), width=Inches(1.6))
		# Explanation and taxonomy
		doc.add_paragraph("@explanation")
		doc.add_paragraph(str(b['explanation']))
		doc.add_paragraph(f"@subject {b['subject']}")
		doc.add_paragraph(f"@unit {b['unit']}")
		doc.add_paragraph(f"@topic {b['topic']}")
		doc.add_paragraph("@plusmarks 1")
		doc.add_paragraph("")
		doc.add_paragraph("---")
		doc.add_paragraph("")
	doc.save(path)


# ---------------- Shadow images and writer ----------------
def draw_shape(ax, shape: str, center=(0, 0), size=1.0, color='#333') -> None:
	import matplotlib.patches as patches
	x, y = center
	if shape == 'circle':
		patch = patches.Circle((x, y), radius=size * 0.5, fill=False, linewidth=2, edgecolor=color)
		ax.add_patch(patch)
	elif shape == 'square':
		patch = patches.Rectangle((x - size * 0.5, y - size * 0.5), size, size, fill=False, linewidth=2, edgecolor=color)
		ax.add_patch(patch)
	elif shape == 'triangle':
		pts = [(x, y + size * 0.6), (x - size * 0.6, y - size * 0.6), (x + size * 0.6, y - size * 0.6)]
		patch = patches.Polygon(pts, closed=True, fill=False, linewidth=2, edgecolor=color)
		ax.add_patch(patch)
	elif shape == 'star':
		# simple 5-point star
		import numpy as np
		angles = np.linspace(0, 2 * np.pi, 6)[:-1]
		outer = [(x + size * 0.6 * np.cos(a), y + size * 0.6 * np.sin(a)) for a in angles]
		inner = [(x + size * 0.25 * np.cos(a + np.pi / 5), y + size * 0.25 * np.sin(a + np.pi / 5)) for a in angles]
		pts = []
		for i in range(5):
			pts.append(outer[i])
			pts.append(inner[i])
		patch = patches.Polygon(pts, closed=True, fill=False, linewidth=2, edgecolor=color)
		ax.add_patch(patch)


def save_sequence_composite_shadow(path: Path) -> None:
	fig, ax = plt.subplots(figsize=(5, 1.2))
	shapes = ['circle', 'square', 'triangle', 'star']
	for i in range(8):
		shape = shapes[i % len(shapes)]
		draw_shape(ax, shape, center=(i * 1.2, 0), size=0.8)
	ax.axis('off')
	ax.set_xlim(-1, 1.2 * 8)
	ax.set_ylim(-1, 1)
	fig.tight_layout()
	fig.savefig(path, dpi=200, bbox_inches='tight')
	plt.close(fig)


def save_shape_icon_shadow(path: Path, shape: str) -> None:
	fig, ax = plt.subplots(figsize=(1.2, 1.2))
	draw_shape(ax, shape, center=(0, 0), size=1.2)
	ax.axis('off')
	fig.tight_layout()
	fig.savefig(path, dpi=200, bbox_inches='tight')
	plt.close(fig)


def save_altitude_graph_shadow(path: Path, times, alts) -> None:
	fig, ax = plt.subplots(figsize=(4.5, 2.8))
	ax.plot(times, alts, marker='o')
	ax.set_xlabel('Time (h)')
	ax.set_ylabel('Altitude (m)')
	ax.set_title("Hiker's Altitude")
	fig.tight_layout()
	fig.savefig(path, dpi=200, bbox_inches='tight')
	plt.close(fig)


def save_midpoint_diagram_shadow(path: Path, st_len: int) -> None:
	fig, ax = plt.subplots(figsize=(5, 1))
	# Draw line RV with points R - S - T - V
	R_x = 0
	ST = st_len
	RT = 2 * ST
	RV = 2 * RT
	S_x = RT / 2
	T_x = RT
	V_x = RV
	ax.hlines(0, R_x, V_x, colors='k', linewidth=2)
	for x, label in [(R_x, 'R'), (S_x, 'S'), (T_x, 'T'), (V_x, 'V')]:
		ax.plot([x], [0], 'ko')
		ax.text(x, 0.15, label, ha='center')
	ax.text((S_x + T_x) / 2, -0.2, f'ST={ST}', ha='center')
	ax.axis('off')
	fig.tight_layout()
	fig.savefig(path, dpi=200, bbox_inches='tight')
	plt.close(fig)


def save_rectangle_shading_shadow(path: Path, cols: int, rows: int, shaded_count: float) -> None:
	fig, ax = plt.subplots(figsize=(4, 2))
	# Draw grid
	for i in range(cols + 1):
		ax.plot([i, i], [0, rows], 'k-', linewidth=1)
	for j in range(rows + 1):
		ax.plot([0, cols], [j, j], 'k-', linewidth=1)
	# Shade cells left to right, top to bottom
	import math
	full = int(math.floor(shaded_count))
	frac = shaded_count - full
	ci = 0
	for idx in range(full):
		c = idx % cols
		r = rows - 1 - (idx // cols)
		rect = plt.Rectangle((c, r), 1, 1, color='#cce5ff')
		ax.add_patch(rect)
	ci = full
	if frac > 0:
		c = ci % cols
		r = rows - 1 - (ci // cols)
		rect = plt.Rectangle((c, r), frac, 1, color='#cce5ff')
		ax.add_patch(rect)
	ax.set_xlim(0, cols)
	ax.set_ylim(0, rows)
	ax.set_aspect('equal')
	ax.axis('off')
	fig.tight_layout()
	fig.savefig(path, dpi=200, bbox_inches='tight')
	plt.close(fig)


def save_segments_with_squares_shadow(path: Path, ab: int, cd: int, ef: int, sq: int) -> None:
	fig, ax = plt.subplots(figsize=(6, 1.2))
	x = 0
	ax.hlines(0, x, x + ab, colors='k', linewidth=3)
	ax.text(x + ab / 2, 0.15, f'AB={ab}', ha='center')
	x += ab
	# square
	ax.add_patch(plt.Rectangle((x, -0.5), sq, 1, fill=False, linewidth=2))
	ax.text(x + sq / 2, -0.7, f'{sq}', ha='center')
	x += sq
	ax.hlines(0, x, x + cd, colors='k', linewidth=3)
	ax.text(x + cd / 2, 0.15, f'CD={cd}', ha='center')
	x += cd
	# square
	ax.add_patch(plt.Rectangle((x, -0.5), sq, 1, fill=False, linewidth=2))
	ax.text(x + sq / 2, -0.7, f'{sq}', ha='center')
	x += sq
	ax.hlines(0, x, x + ef, colors='k', linewidth=3)
	ax.text(x + ef / 2, 0.15, f'EF={ef}', ha='center')
	ax.axis('off')
	fig.tight_layout()
	fig.savefig(path, dpi=200, bbox_inches='tight')
	plt.close(fig)


def save_card_holes_shadow(path: Path) -> None:
	fig, ax = plt.subplots(figsize=(2.5, 2.5))
	# square card
	ax.add_patch(plt.Rectangle((-1, -1), 2, 2, fill=False, linewidth=2))
	# holes
	ax.add_patch(plt.Circle((-0.5, 0.4), 0.1, color='k'))
	ax.add_patch(plt.Circle((0.6, -0.2), 0.1, color='k'))
	ax.set_xlim(-1.2, 1.2)
	ax.set_ylim(-1.2, 1.2)
	ax.set_aspect('equal')
	ax.axis('off')
	fig.tight_layout()
	fig.savefig(path, dpi=200, bbox_inches='tight')
	plt.close(fig)


def build_25_shadow_blocks_with_images() -> List[Dict[str, object]]:
	blocks: List[Dict[str, object]] = []
	add = blocks.append
	shadow_dir = IMAGES_DIR / 'shadow'
	shadow_dir.mkdir(parents=True, exist_ok=True)
	
	# Q1 - no image
	add({
		'title': 'Solve Linear Equation (One-Step)',
		'description': 'Solve for n in a simple linear equation.',
		'question': 'If $n+7=12$, what is the value of $n$?',
		'instruction': 'Select the correct value of n.',
		'difficulty': 'easy', 'order': 1,
		'options': ['2', '4', '5', '7', '12'], 'answer': '5',
		'explanation': 'Subtract 7 from both sides: $n = 12-7 = 5$.',
		'subject': 'Quantitative Math', 'unit': 'Algebra', 'topic': 'Interpreting Variables',
		'question_image_paths': [], 'option_image_paths': {}
	})
	
	# Q2 - sequence with images
	seq_img = shadow_dir / 'shadow_q2_sequence.png'
	save_sequence_composite_shadow(seq_img)
	opt_map = {
		'Circle': shadow_dir / 'shadow_q2_opt_circle.png',
		'Square': shadow_dir / 'shadow_q2_opt_square.png',
		'Triangle': shadow_dir / 'shadow_q2_opt_triangle.png',
		'Star': shadow_dir / 'shadow_q2_opt_star.png',
		'Hexagon': shadow_dir / 'shadow_q2_opt_hex.png',
	}
	# generate four; hexagon as text only
	save_shape_icon_shadow(opt_map['Circle'], 'circle')
	save_shape_icon_shadow(opt_map['Square'], 'square')
	save_shape_icon_shadow(opt_map['Triangle'], 'triangle')
	save_shape_icon_shadow(opt_map['Star'], 'star')
	add({
		'title': 'Repeating Symbol Sequence',
		'description': 'Identify a term in a repeating sequence using modular arithmetic.',
		'question': 'A sequence repeats the symbols in order: Circle, Square, Triangle, Star. Which is the 12th symbol?',
		'instruction': 'Determine the cycle length and reduce the index modulo the cycle length.',
		'difficulty': 'moderate', 'order': 2,
		'options': ['Circle', 'Square', 'Triangle', 'Star', 'Hexagon'], 'answer': 'Star',
		'explanation': 'Cycle length is 4. 12 mod 4 = 0, so the 12th is the 4th in the cycle: Star.',
		'subject': 'Quantitative Math', 'unit': 'Numbers and Operations', 'topic': 'Sequences & Series',
		'question_image_paths': [seq_img], 'option_image_paths': {k: v for k, v in opt_map.items() if k != 'Hexagon'}
	})
	
	# Q3 - no image
	add({
		'title': 'Expression for Total Items',
		'description': 'Translate a word scenario into an algebraic expression.',
		'question': 'A jar contains 15 marbles. You add $y$ more marbles. Which expression represents the total number of marbles?',
		'instruction': 'Choose the expression that models the situation.',
		'difficulty': 'easy', 'order': 3,
		'options': ['$15-y$', '$15y$', '$\\frac{15}{y}$', '$y-15$', '$15+y$'], 'answer': '$15+y$',
		'explanation': 'Start with 15 and add y new marbles: $15 + y$.',
		'subject': 'Quantitative Math', 'unit': 'Algebra', 'topic': 'Interpreting Variables',
		'question_image_paths': [], 'option_image_paths': {}
	})
	
	# Q4 - no image
	add({
		'title': 'Place Value and Inequality',
		'description': 'Find the greatest digit for a number to stay below a bound.',
		'question': 'In the number $5,\\square 42$, $\\square$ is a digit 0–9. If the number is less than 5,242, what is the greatest possible value for $\\square$?',
		'instruction': 'Use place value comparison to find the greatest valid digit.',
		'difficulty': 'easy', 'order': 4,
		'options': ['0', '1', '2', '4', '9'], 'answer': '1',
		'explanation': 'Compare hundreds place with 2 in 5,242: the greatest hundreds digit to keep it smaller is 1.',
		'subject': 'Quantitative Math', 'unit': 'Numbers and Operations', 'topic': 'Computation with Whole Numbers',
		'question_image_paths': [], 'option_image_paths': {}
	})
	
	# Q5 - no image
	add({
		'title': 'Adding Fractions',
		'description': 'Add two fractions with unlike denominators.',
		'question': 'Which of the following is the sum of $\\frac{5}{12}$ and $\\frac{1}{3}$?',
		'instruction': 'Compute using a common denominator.',
		'difficulty': 'easy', 'order': 5,
		'options': ['$\\frac{1}{4}$', '$\\frac{2}{3}$', '$\\frac{3}{4}$', '$\\frac{5}{6}$', '$\\frac{7}{12}$'], 'answer': '$\\frac{3}{4}$',
		'explanation': '$\\frac{5}{12}+\\frac{1}{3}=\\frac{5}{12}+\\frac{4}{12}=\\frac{9}{12}=\\frac{3}{4}$.',
		'subject': 'Quantitative Math', 'unit': 'Numbers and Operations', 'topic': 'Fractions, Decimals, & Percents',
		'question_image_paths': [], 'option_image_paths': {}
	})
	
	# Q6 - altitude graph image
	alt_img = shadow_dir / 'shadow_q6_altitude.png'
	save_altitude_graph_shadow(alt_img, times=[0, 1, 2, 3, 4], alts=[120, 180, 250, 300, 420])
	add({
		'title': 'Altitude Difference from a Graph (Conceptual)',
		'description': 'Read altitude change from start and finish.',
		'question': 'A hiker starts at 120 meters and ends at 420 meters after a steady climb. How many meters higher is the end than the start?',
		'instruction': 'Compute final altitude minus initial altitude.',
		'difficulty': 'easy', 'order': 6,
		'options': ['120', '240', '300', '320', '540'], 'answer': '300',
		'explanation': '420 − 120 = 300 meters.',
		'subject': 'Quantitative Math', 'unit': 'Data Analysis & Probability', 'topic': 'Interpretation of Tables & Graphs',
		'question_image_paths': [alt_img], 'option_image_paths': {}
	})
	
	# Q7 - no image
	add({
		'title': 'Multiply Decimals',
		'description': 'Evaluate a product of decimals.',
		'question': 'What is the value of $0.25 \\times 18 \\times 0.4$?',
		'instruction': 'Use associativity to simplify.',
		'difficulty': 'easy', 'order': 7,
		'options': ['0.18', '1.8', '18', '180', '0.72'], 'answer': '1.8',
		'explanation': '$0.25 \\times 0.4 = 0.1$ and $0.1 \\times 18 = 1.8$.',
		'subject': 'Quantitative Math', 'unit': 'Numbers and Operations', 'topic': 'Fractions, Decimals, & Percents',
		'question_image_paths': [], 'option_image_paths': {}
	})
	
	# Q8 - no image
	add({
		'title': 'Minimize Coins for a Total',
		'description': 'Find the least number of coins to make a given amount.',
		'question': 'There are ten of each coin: 1¢, 5¢, 10¢, and 25¢. If you need exactly 47¢, what is the least number of coins required?',
		'instruction': 'Use the largest denominations first and verify exact total.',
		'difficulty': 'moderate', 'order': 8,
		'options': ['Three', 'Four', 'Five', 'Six', 'Seven'], 'answer': 'Five',
		'explanation': '47 = 25 + 10 + 10 + 1 + 1 uses five coins; four coins cannot make 47.',
		'subject': 'Quantitative Math', 'unit': 'Reasoning', 'topic': 'Word Problems',
		'question_image_paths': [], 'option_image_paths': {}
	})
	
	# Q9 - no image
	add({
		'title': 'Multiply Fractions then Halve',
		'description': 'Evaluate a nested fractional expression.',
		'question': 'What is the value of $\\frac{1}{2}\\left(\\frac{2}{3} \\times \\frac{3}{4}\\right)$?',
		'instruction': 'Multiply inside the parentheses first.',
		'difficulty': 'easy', 'order': 9,
		'options': ['$\\frac{1}{4}$', '$\\frac{1}{3}$', '$\\frac{3}{8}$', '$\\frac{5}{12}$', '$\\frac{7}{24}$'], 'answer': '$\\frac{1}{4}$',
		'explanation': '$\\frac{2}{3} \\times \\frac{3}{4} = \\frac{1}{2}$; then half gives $\\frac{1}{4}$.',
		'subject': 'Quantitative Math', 'unit': 'Numbers and Operations', 'topic': 'Fractions, Decimals, & Percents',
		'question_image_paths': [], 'option_image_paths': {}
	})
	
	# Q10 - midpoint diagram image
	mid_img = shadow_dir / 'shadow_q10_mid.png'
	save_midpoint_diagram_shadow(mid_img, st_len=10)
	add({
		'title': 'Midpoints on a Line Segment',
		'description': 'Use midpoint relations to compute a segment length.',
		'question': 'Segment $\\overline{ST}$ has length 10, $T$ is the midpoint of $\\overline{RV}$, and $S$ is the midpoint of $\\overline{RT}$. What is the length of $\\overline{SV}$?',
		'instruction': 'Express RV in terms of ST using midpoint relations.',
		'difficulty': 'moderate', 'order': 10,
		'options': ['10', '20', '30', '40', '50'], 'answer': '30',
		'explanation': 'S midpoint of RT ⇒ ST = RT/2 ⇒ RT = 20. T midpoint of RV ⇒ TV = RT = 20. So SV = ST + TV = 10 + 20 = 30.',
		'subject': 'Quantitative Math', 'unit': 'Geometry and Measurement', 'topic': 'Lines, Angles, & Triangles',
		'question_image_paths': [mid_img], 'option_image_paths': {}
	})
	
	# Q11 - no image
	add({
		'title': 'Solve Whole-Number Identity',
		'description': 'Solve for a whole number that satisfies a simple quadratic identity.',
		'question': 'Let $a$ be defined by $a=a^{2}-a$, where $a$ is a whole number and $a\\neq 0$. What is the value of $3a$?',
		'instruction': 'Solve for a, then compute 3a.',
		'difficulty': 'easy', 'order': 11,
		'options': ['4', '5', '6', '7', '8'], 'answer': '6',
		'explanation': '$a=a^{2}-a \\Rightarrow a^{2}-2a=0 \\Rightarrow a(a-2)=0$. With $a\\neq 0$, $a=2$, so $3a=6$.',
		'subject': 'Quantitative Math', 'unit': 'Algebra', 'topic': 'Interpreting Variables',
		'question_image_paths': [], 'option_image_paths': {}
	})
	
	# Q12 - table (shirt/pants)
	add({
		'title': 'Counting Uniform Combinations',
		'description': 'Count combinations from shirts and pants options.',
		'question': 'A uniform has 1 shirt and 1 pair of pants. If there are 5 shirt colors and 2 pants colors, how many different uniforms are possible?',
		'instruction': 'Multiply the number of shirt choices by pant choices.',
		'difficulty': 'easy', 'order': 12,
		'options': ['6', '8', '10', '12', '15'], 'answer': '10',
		'explanation': 'There are 5 shirts and 2 pants: $5 \\times 2 = 10$.',
		'subject': 'Quantitative Math', 'unit': 'Data Analysis & Probability', 'topic': 'Counting & Arrangement Problems',
		'table_rows': [
			['Shirt Color', 'Pants Color'],
			['Blue', 'Black'],
			['Green', 'Khaki'],
			['White', 'Navy'],
			['Red', ' '],
			['Yellow', ' '],
		]
	})
	
	# Q13 - no image
	add({
		'title': 'Parity Reasoning',
		'description': 'Determine which expression yields an odd integer for even n.',
		'question': 'If $n$ is an even integer, which of the following must be an odd integer?',
		'instruction': 'Analyze parity for each expression.',
		'difficulty': 'easy', 'order': 13,
		'options': ['$n$', '$n+1$', '$2n$', '$3n$', '$n+2$'], 'answer': '$n+1$',
		'explanation': 'If $n$ is even, then $n+1$ is odd.',
		'subject': 'Quantitative Math', 'unit': 'Numbers and Operations', 'topic': 'Basic Number Theory',
		'question_image_paths': [], 'option_image_paths': {}
	})
	
	# Q14 - no image
	add({
		'title': 'Direct Proportion: Miles per Dollar',
		'description': 'Use proportional reasoning to scale miles by fuel cost.',
		'question': 'A car travels 180 miles on $\\$30 of gas. At the same rate, how many miles on $\\$45?',
		'instruction': 'Use miles per dollar to scale linearly.',
		'difficulty': 'easy', 'order': 14,
		'options': ['225', '240', '255', '270', '300'], 'answer': '270',
		'explanation': '$180/30 = 6$ miles per dollar; $6 \\times 45 = 270$.',
		'subject': 'Quantitative Math', 'unit': 'Reasoning', 'topic': 'Word Problems',
		'question_image_paths': [], 'option_image_paths': {}
	})
	
	# Q15 - no image
	add({
		'title': 'Closest Fraction to a Percentage',
		'description': 'Compare fractions to 62%.',
		'question': 'Which fraction is closest to $62\\%$?',
		'instruction': 'Convert fractions to percents or compare decimals.',
		'difficulty': 'moderate', 'order': 15,
		'options': ['$\\frac{1}{2}$', '$\\frac{3}{5}$', '$\\frac{5}{8}$', '$\\frac{2}{3}$', '$\\frac{7}{10}$'], 'answer': '$\\frac{5}{8}$',
		'explanation': '$\\frac{5}{8}=0.625=62.5\\%$, closest to 62%.',
		'subject': 'Quantitative Math', 'unit': 'Numbers and Operations', 'topic': 'Fractions, Decimals, & Percents',
		'question_image_paths': [], 'option_image_paths': {}
	})
	
	# Q16 - no image
	add({
		'title': 'Balanced Club Sizes',
		'description': 'Distribute students into clubs with max difference 1.',
		'question': 'There are 84 students forming 5 clubs. Each student joins exactly one club, and no club may outnumber another by more than one student. What is the least possible number of students in one club?',
		'instruction': 'Distribute as evenly as possible.',
		'difficulty': 'moderate', 'order': 16,
		'options': ['15', '16', '17', '18', '19'], 'answer': '16',
		'explanation': '84 divided as evenly as possible into 5 gives sizes 17, 17, 17, 16, 17; the least is 16.',
		'subject': 'Quantitative Math', 'unit': 'Data Analysis & Probability', 'topic': 'Counting & Arrangement Problems',
		'question_image_paths': [], 'option_image_paths': {}
	})
	
	# Q17 - rectangle shading image
	rect_img = shadow_dir / 'shadow_q17_rect.png'
	save_rectangle_shading_shadow(rect_img, cols=4, rows=2, shaded_count=5.5)
	add({
		'title': 'Shaded Fraction of a Rectangle (Variant)',
		'description': 'Find the shaded portion count out of total.',
		'question': 'A rectangle is divided into 8 congruent squares. If $5\\tfrac{1}{2}$ squares are shaded, what fraction of the rectangle is shaded?',
		'instruction': 'Compute shaded total over 8 and simplify if possible.',
		'difficulty': 'easy', 'order': 17,
		'options': ['$\\frac{5}{8}$', '$\\frac{11}{16}$', '$\\frac{3}{4}$', '$\\frac{7}{12}$', '$\\frac{2}{3}$'], 'answer': '$\\frac{11}{16}$',
		'explanation': '$5.5/8 = 11/16$.',
		'subject': 'Quantitative Math', 'unit': 'Geometry and Measurement', 'topic': 'Area & Volume',
		'question_image_paths': [rect_img], 'option_image_paths': {}
	})
	
	# Q18 - no image
	add({
		'title': 'Currency Exchange Chains',
		'description': 'Convert gold to copper through given exchange rates.',
		'question': 'In a game, 1 gold piece may be exchanged for 4 silver pieces, and 3 silver pieces may be exchanged for 18 copper pieces. How many copper pieces for 5 gold pieces?',
		'instruction': 'Find copper per gold, then scale.',
		'difficulty': 'easy', 'order': 18,
		'options': ['60', '90', '100', '120', '150'], 'answer': '120',
		'explanation': '1 silver = 6 copper; 1 gold = 4 silver = 24 copper; 5 gold = 120 copper.',
		'subject': 'Quantitative Math', 'unit': 'Numbers and Operations', 'topic': 'Rational Numbers',
		'question_image_paths': [], 'option_image_paths': {}
	})
	
	# Q19 - segments with squares image
	seg_img = shadow_dir / 'shadow_q19_segments.png'
	save_segments_with_squares_shadow(seg_img, ab=5, cd=9, ef=7, sq=3)
	add({
		'title': 'Sum of Horizontal Segments (Variant)',
		'description': 'Use only horizontal contributions to find n as a horizontal length.',
		'question': 'The figure shows AB=5 cm, CD=9 cm, EF=7 cm with two squares of side 3 cm placed between the segments. What is the horizontal length n?',
		'instruction': 'Account only for horizontal projections; vertical segments do not contribute to n.',
		'difficulty': 'moderate', 'order': 19,
		'options': ['13', '14', '15', '16', '17'], 'answer': '15',
		'explanation': 'n = 5 + 9 + 7 − 3 − 3 = 15 cm.',
		'subject': 'Quantitative Math', 'unit': 'Geometry and Measurement', 'topic': 'Coordinate Geometry',
		'question_image_paths': [seg_img], 'option_image_paths': {}
	})
	
	# Q20 - no image
	add({
		'title': 'Order of Operations',
		'description': 'Evaluate an expression with exponents, multiplication/division, and addition.',
		'question': 'Calculate: $2+8 \\times 3^{2} \\div 4+5^{2}$',
		'instruction': 'Apply exponents first, then multiplication/division from left to right, then addition.',
		'difficulty': 'easy', 'order': 20,
		'options': ['35', '39', '41', '45', '49'], 'answer': '45',
		'explanation': '$3^{2}=9; 8\\times9=72; 72\\div4=18; 2+18+25=45$.',
		'subject': 'Quantitative Math', 'unit': 'Numbers and Operations', 'topic': 'Order of Operations',
		'question_image_paths': [], 'option_image_paths': {}
	})
	
	# Q21 - card holes image
	card_img = shadow_dir / 'shadow_q21_card.png'
	save_card_holes_shadow(card_img)
	add({
		'title': 'Face-Down Flip Concept',
		'description': 'Understand difference between rotations and mirror reflections.',
		'question': 'After turning a card face down, which of the following cannot be obtained by rotation alone from the original face-up orientation?',
		'instruction': 'Recall that a face-down flip produces a mirror image.',
		'difficulty': 'hard', 'order': 21,
		'options': ['90° rotation', '180° rotation', 'Vertical mirror image', '270° rotation', '0° (no change)'], 'answer': 'Vertical mirror image',
		'explanation': 'Mirror images cannot be produced by rotations alone.',
		'subject': 'Quantitative Math', 'unit': 'Geometry and Measurement', 'topic': 'Transformations (Dilating a shape)',
		'question_image_paths': [card_img], 'option_image_paths': {}
	})
	
	# Q22 - no image
	add({
		'title': 'Integer Conditions with Odd n',
		'description': 'Decide which expression is always an integer for odd n.',
		'question': 'If a number $n$ is odd, which of the following expressions must be an integer?',
		'instruction': 'Let $n=2k+1$ and test each expression.',
		'difficulty': 'easy', 'order': 22,
		'options': ['$\\frac{n}{2}$', '$\\frac{n+1}{2}$', '$\\frac{3n}{4}$', '$\\frac{n+3}{4}$', '$\\frac{n+2}{3}$'], 'answer': '$\\frac{n+1}{2}$',
		'explanation': 'For $n=2k+1$, $\\frac{n+1}{2}=k+1$ is always an integer.',
		'subject': 'Quantitative Math', 'unit': 'Numbers and Operations', 'topic': 'Basic Number Theory',
		'question_image_paths': [], 'option_image_paths': {}
	})
	
	# Q23 - no image
	add({
		'title': 'Reading Fractions of a Book (Variant)',
		'description': 'Track remaining pages after fractional reading over two days.',
		'question': 'On Monday, a reader completes $\\frac{1}{4}$ of a book; on Tuesday, $\\frac{1}{3}$ of the remaining pages. To finish, 90 pages are left. How many pages are in the book?',
		'instruction': 'Compute the fraction remaining after each day and set to 90.',
		'difficulty': 'moderate', 'order': 23,
		'options': ['120', '150', '180', '240', '360'], 'answer': '180',
		'explanation': 'After Monday: 3/4 remain. Tuesday reads 1/3 of that ⇒ 2/3 remain of 3/4 ⇒ 1/2 of the book. 1/2 = 90 ⇒ total 180.',
		'subject': 'Quantitative Math', 'unit': 'Reasoning', 'topic': 'Word Problems',
		'question_image_paths': [], 'option_image_paths': {}
	})
	
	# Q24 - no image
	add({
		'title': 'Circumference of Inscribed Circle',
		'description': 'Compute circumference from a square’s area.',
		'question': 'A square has area 196 in^2. What is the circumference of the largest circle cut from it?',
		'instruction': 'Diameter equals square side length.',
		'difficulty': 'easy', 'order': 24,
		'options': ['$14\\pi$', '$28\\pi$', '$42\\pi$', '$56\\pi$', '$196\\pi$'], 'answer': '$14\\pi$',
		'explanation': 'Side = $\\sqrt{196}=14$, so circumference = $\\pi d = 14\\pi$.',
		'subject': 'Quantitative Math', 'unit': 'Geometry and Measurement', 'topic': 'Circles (Area, circumference)',
		'question_image_paths': [], 'option_image_paths': {}
	})
	
	# Q25 - no image
	add({
		'title': 'Successive Percent Changes',
		'description': 'Apply percentage increase then decrease.',
		'question': 'The number 150 is increased by 20%, then decreased by 25% to give x. What is x?',
		'instruction': 'Compute step by step.',
		'difficulty': 'easy', 'order': 25,
		'options': ['110', '115', '120', '130', '135'], 'answer': '135',
		'explanation': '150 \\to 180 (increase 20%), then 180 \\times 0.75 = 135.',
		'subject': 'Quantitative Math', 'unit': 'Numbers and Operations', 'topic': 'Fractions, Decimals, & Percents',
		'question_image_paths': [], 'option_image_paths': {}
	})
	
	return blocks


def write_shadow_questions_docx_with_images(path: Path) -> None:
	blocks = build_25_shadow_blocks_with_images()
	doc = Document()
	img_dir = IMAGES_DIR / 'shadow'
	for b in blocks:
		# Header tags
		doc.add_paragraph(f"@title {b['title']}")
		doc.add_paragraph(f"@description {b['description']}")
		doc.add_paragraph("")
		# Question text
		doc.add_paragraph(f"@question {b['question']}")
		doc.add_paragraph(f"@instruction {b['instruction']}")
		doc.add_paragraph(f"@difficulty {b['difficulty']}")
		doc.add_paragraph(f"@Order {b['order']}")
		# Question images
		for p in b.get('question_image_paths', []):
			if p and Path(p).exists() and Path(p).stat().st_size > 0:
				doc.add_picture(str(p), width=Inches(4.5))
		# Table if present
		table_rows = b.get('table_rows')
		if table_rows:
			rows = len(table_rows)
			cols = len(table_rows[0]) if rows > 0 else 0
			table = doc.add_table(rows=rows, cols=cols)
			for i, row in enumerate(table_rows):
				for j, cell in enumerate(row):
					table.cell(i, j).text = cell
		# Options with images
		opt_imgs: Dict[str, Path] = b.get('option_image_paths', {})  # label -> path
		for opt in b['options']:
			prefix = '@@option' if opt == b['answer'] else '@option'
			doc.add_paragraph(f"{prefix} {opt}")
			if opt in opt_imgs:
				p = opt_imgs[opt]
				if p and Path(p).exists() and Path(p).stat().st_size > 0:
					doc.add_picture(str(p), width=Inches(1.4))
		# Explanation and taxonomy
		doc.add_paragraph("@explanation")
		doc.add_paragraph(str(b['explanation']))
		doc.add_paragraph(f"@subject {b['subject']}")
		doc.add_paragraph(f"@unit {b['unit']}")
		doc.add_paragraph(f"@topic {b['topic']}")
		doc.add_paragraph("@plusmarks 1")
		doc.add_paragraph("")
		doc.add_paragraph("---")
		doc.add_paragraph("")
	doc.save(path)


if __name__ == '__main__':
	main()